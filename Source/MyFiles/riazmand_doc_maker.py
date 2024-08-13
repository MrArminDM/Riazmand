#region library
import datetime, os
from docx import Document               #for creating and manipulating Word documents
from docx.shared import Pt, RGBColor            #Change lines font size and color
from docx.enum.text import WD_ALIGN_PARAGRAPH   #Make lines align cneter
#endregion library

#region Global
logger_mode = "on"
#region Doc
doc = ""
jock_doc = Document()        #easier read & write
log_doc = Document()         #easier read & write
#endregion Doc
#endregion Global
#region Main
#region DateTime
def get_datetime(mode):
    now = datetime.datetime.now()
    if mode == "current_time":
        return now.strftime("%Y/%m/%d - %H:%M:%S")
    elif mode == "file_name":
        return now.strftime("%Y.%m.%d_%H.%M.%S")
#endregion DateTime
#region Logger
def log_mode(mode):
    global logger_mode
    if mode == "on":
        logger_mode = "on"
    else:
        logger_mode = "off"
#endregion Logger
#region Doc
#region Main

def jock_to_file(language, text, doc_name):  #to export jokes to a Word document
    global doc
    doc = jock_doc
    if language == "En":
        header("Armin Dost Mohammadi", "Riazmand", "Extract jock")
    else:
        header("استخراج جوک", "ریاضمند", "آرمین دوست محمدی")

    normal_format(f"{text}\n" + "-" * 26)

    doc_saver(doc_name)

def logger(language, log):  #to log events to a Word document
    global doc
    doc = log_doc
    if logger_mode == "on":
        if language == "En":
            header("Armin Dost Mohammadi", "Riazmand", "Log")
        else:
            header("گزارش", "ریاضمند", "آرمین دوست محمدی")

        if language == "Fa":
            text = log[0]
        else:
            text = log[1]

        normal_format(f"{get_datetime("current_time")}\n{text}\n" + "-" * 26)
#endregion Main
#region Tool

def doc_title(title, mode):
    global doc
    if mode == "log":
        doc = log_doc
    else:
        doc = jock_doc

    title_format(f"{title}\n")


def doc_saver(fname):
    folder_path = 'C:\\Riazmand/Exports/'
    if not os.path.exists('C:\\Riazmand'):
        os.makedirs("C:\\Riazmand")
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    file_path = f'C:\\Riazmand/Exports/{fname}'
    if os.path.exists(file_path):
        os.remove(file_path)
    doc.save(f'C:\\Riazmand/Exports/{fname}')

def log_saver():
    doc = log_doc
    folder_path = 'C:\\Riazmand/Log/'
    if not os.path.exists('C:\\Riazmand'):
        os.makedirs("C:\\Riazmand")
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    
    doc.save(f"C:\\Riazmand/Log/riazmand.log_{get_datetime("file_name")}.docx")

def clear_doc():
    global jock_doc
    jock_doc = ""
    jock_doc = Document()

def header(left, center, right):  #to format text in the Word document
    global doc
    header = doc.sections[0].header.paragraphs[0]
    header.text = f"{left}\t{center}\t{right}"
    header.style = doc.styles["Header"]

def title_format(text):  #to format text in the Word document
    global doc
    pr = doc.add_paragraph()
    run = pr.add_run(text)

    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(110, 173, 255)
    pr.paragraph_format.keep_together = True
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER

def normal_format(text):  #to format text in the Word document
    global doc
    pr = doc.add_paragraph()
    run = pr.add_run(text)

    run.font.name = 'Arial'
    run.font.size = Pt(12)
    pr.paragraph_format.keep_together = True
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
#endregion Tool
#endregion Doc
#endregion Main

doc_title(f"Start date: {get_datetime("current_time")}", "log")